"""
REMIS Backend Report – Part 2  (Sections 9-16)
Appends to the existing REMIS_Backend_Report.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"c:\xampp\htdocs\RemisV2"
IMGS = os.path.join(BASE, "docs", "images")
PUB  = os.path.join(BASE, "public", "images")
OUT  = os.path.join(BASE, "docs", "REMIS_Backend_Report.docx")

doc = Document(OUT)

# ── helpers (same as part 1) ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),hex_color); tcPr.append(shd)

def add_page_break(doc): doc.add_page_break()

def heading(doc,text,level=1,color="1e3a5f"):
    p=doc.add_heading(text,level=level)
    for r in p.runs: r.font.color.rgb=RGBColor.from_string(color)
    return p

def subheading(doc,text,color="2e6da4"): return heading(doc,text,2,color)
def sub3(doc,text,color="17a2b8"):       return heading(doc,text,3,color)

def para(doc,text="",bold=False,italic=False,size=11,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    return p

def code_block(doc,text,size=8):
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),"F3F4F6"); pPr.append(shd)
    ind=OxmlElement("w:ind")
    ind.set(qn("w:left"),"360"); ind.set(qn("w:right"),"360"); pPr.append(ind)
    r=p.add_run(text); r.font.name="Courier New"; r.font.size=Pt(size)
    r.font.color.rgb=RGBColor(0x1e,0x3a,0x5f)
    return p

def bullet(doc,text,level=0):
    p=doc.add_paragraph(text,style="List Bullet")
    p.paragraph_format.left_indent=Inches(0.25*(level+1))
    return p

def insert_image(doc,path,width=6.0,caption=None):
    if not os.path.exists(path):
        para(doc,f"[Image not found: {path}]",italic=True,color="999999"); return
    doc.add_picture(path,width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size=Pt(9); r.font.italic=True
            r.font.color.rgb=RGBColor(0x66,0x66,0x66)

def add_table(doc,headers,rows,col_widths=None):
    tbl=doc.add_table(rows=1,cols=len(headers))
    tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=tbl.rows[0].cells
    for i,h in enumerate(headers):
        hc[i].text=h; set_cell_bg(hc[i],"1e3a5f")
        r=hc[i].paragraphs[0].runs[0]
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.bold=True; r.font.size=Pt(9)
        hc[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,row in enumerate(rows):
        cells=tbl.add_row().cells
        bg="F8F9FF" if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(row):
            cells[ci].text=str(val); set_cell_bg(cells[ci],bg)
            cells[ci].paragraphs[0].runs[0].font.size=Pt(9)
    if col_widths:
        for row in tbl.rows:
            for ci,w in enumerate(col_widths): row.cells[ci].width=Inches(w)
    return tbl

# ════════════════════════════════════════════════════════════════════════════
# 9. MIDDLEWARE
# ════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
heading(doc,"9. Middleware")
para(doc,
    "Three custom middleware classes protect routes and enforce business rules. "
    "All three are registered as aliases in bootstrap/app.php and applied to "
    "route groups in routes/web.php.")

add_table(doc,
    ["Alias","Class","Applied To","What It Does"],
    [
        ["role",                  "EnsureRole",           "landlord/tenant/admin route groups",  "abort(403) if user.role != required role"],
        ["locked",                "CheckSessionLock",     "landlord/tenant/admin route groups",  "Redirects to /lock for remembered or locked sessions"],
        ["force.password.change", "ForcePasswordChange",  "tenant route group only",             "Redirects to /tenant/change-password while flag is true"],
    ],
    col_widths=[1.5,1.8,2.0,2.2]
)

doc.add_paragraph()
subheading(doc,"9.1  EnsureRole  (app/Http/Middleware/EnsureRole.php)")
code_block(doc,"""class EnsureRole
{
    public function handle(Request $request, Closure $next, string $role): mixed
    {
        if (! Auth::check() || Auth::user()->role !== $role) {
            abort(403, 'Unauthorized.');
        }
        return $next($request);
    }
}

// Registered in bootstrap/app.php:
$middleware->alias(['role' => EnsureRole::class]);

// Applied to route groups in routes/web.php:
Route::middleware(['auth', 'role:landlord', 'locked'])->prefix('landlord')->group(...);
Route::middleware(['auth', 'role:tenant',   'force.password.change', 'locked'])->prefix('tenant')->group(...);
Route::middleware(['auth', 'role:admin',    'locked'])->prefix('admin')->group(...);""")

doc.add_paragraph()
subheading(doc,"9.2  CheckSessionLock  (app/Http/Middleware/CheckSessionLock.php)")
para(doc,
    "When Auth::viaRemember() returns true (browser was reopened with a remembered "
    "session), or session('auth.locked') is set, the user is redirected to /lock. "
    "This gives a second-factor check before accessing any protected page.")

doc.add_paragraph()
subheading(doc,"9.3  ForcePasswordChange  (app/Http/Middleware/ForcePasswordChange.php)")
code_block(doc,"""class ForcePasswordChange
{
    public function handle(Request $request, Closure $next): mixed
    {
        $user = Auth::user();

        if ($user && $user->force_password_change) {
            // Only allow the change-password page and logout through
            if (! $request->routeIs('tenant.password.change',
                                     'tenant.password.update', 'logout')) {
                return redirect()->route('tenant.password.change')
                    ->with('warning', 'You must set a new password before continuing.');
            }
        }
        return $next($request);
    }
}""")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 10. ROUTING
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"10. Routing")
para(doc,
    "All 60+ routes are defined in routes/web.php. There is no api.php file — "
    "AJAX endpoints (TIN verification, payment status polling) return JSON from "
    "within the web routes file, sharing the same session-based authentication. "
    "All routes use named route keys for clean URL generation in Blade templates.")

subheading(doc,"10.1  Route Groups Summary")
add_table(doc,
    ["Group","Prefix","Middleware","Route Count"],
    [
        ["Public",   "/",          "(none)",                                       "6"],
        ["Auth Lock","/",          "auth",                                         "2 (GET/POST /lock)"],
        ["Landlord", "/landlord",  "auth, role:landlord, locked",                  "~28"],
        ["Admin",    "/admin",     "auth, role:admin, locked",                     "~10"],
        ["Tenant",   "/tenant",    "auth, role:tenant, force.password.change, locked","~12"],
    ],
    col_widths=[1.3,1.3,2.8,1.0]
)

doc.add_paragraph()
subheading(doc,"10.2  Key Route Definitions")
code_block(doc,"""// ── Landlord: Unit-based lease workflow (new) ──────────────────────────
Route::get('/properties/{property}/units/{unit}/leases/create',
           [LandlordController::class, 'unitLeaseCreate'])
    ->name('properties.units.leases.create');

Route::post('/properties/{property}/units/{unit}/leases',
            [LandlordController::class, 'unitLeaseStore'])
    ->name('properties.units.leases.store');

Route::post('/properties/{property}/units/{unit}/leases/{lease}/assign-tenant',
            [LandlordController::class, 'unitLeaseAssignTenant'])
    ->name('properties.units.leases.assign-tenant');

// ── NMB Payment routes ──────────────────────────────────────────────────
Route::post('/payments/{payment}/generate',
            [PaymentController::class, 'generateControlNumber'])
    ->name('payments.generate');

Route::post('/payments/{payment}/send',
            [PaymentController::class, 'sendControlNumber'])
    ->name('payments.send');

Route::get('/payments/{payment}/status',         // AJAX JSON response
           [PaymentController::class, 'checkStatus'])
    ->name('payments.status');

// ── TIN Verification (AJAX, must be before /{user} wildcard) ───────────
Route::post('/tenants/verify-tin',
            [TinVerificationController::class, 'verify'])
    ->name('tenants.verify-tin');""")

doc.add_paragraph()
subheading(doc,"10.3  Route Naming Convention")
para(doc,
    "All routes follow the pattern {role}.{resource}.{action} — for example "
    "landlord.payments.generate, tenant.settings.profile. This convention "
    "ensures Blade templates and controllers can generate correct URLs with "
    "route('landlord.payments.generate', $payment) regardless of future URL changes.")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 11. CONTROLLERS
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"11. Controllers")
para(doc,
    "Controllers in app/Http/Controllers/ handle all HTTP request/response cycles. "
    "Complex domain logic is delegated to Services, keeping controllers lean.")

add_table(doc,
    ["Controller","Routes Served","Key Responsibilities"],
    [
        ["AuthController",           "/login, /register, /lock, /logout",   "Login, registration with strong validation, session lock/unlock, logout"],
        ["LandlordController",       "/landlord/*",                          "Largest controller (~600+ lines): properties, leases, tenants, reports, settings"],
        ["TenantController",         "/tenant/*",                            "Tenant dashboard, payment checkout, force-password-change, maintenance, settings"],
        ["AdminController",          "/admin/*",                             "System settings, audit logs, backups, user role management"],
        ["PaymentController",        "/landlord/payments/*",                 "NMB control number generation, send email/SMS, status polling, batch poll"],
        ["TinVerificationController","/landlord/tenants/verify-tin",         "AJAX proxy to TRA API with dev-mode mock fallback"],
        ["PasswordSetupController",  "/create-password/{token}",             "Tenant invitation password setup (uses password reset tokens)"],
    ],
    col_widths=[2.0,1.9,2.7]
)

doc.add_paragraph()
subheading(doc,"11.1  LandlordController — Property Creation")
para(doc,
    "The storeProperty() method handles three distinct workflows determined by "
    "the property_category and floor_layout fields submitted from the form.")
code_block(doc,"""// Single-unit property: auto-create one Unit
if ($request->property_category === 'single') {
    $unit = Unit::create([
        'property_id' => $property->id,
        'unit_number' => 'Unit 1',
        'status'      => 'vacant',
    ]);
}

// Multi-unit, single floor: auto-generate "Unit 1", "Unit 2", ...
if ($request->floor_layout === 'single_floor') {
    for ($i = 1; $i <= $request->number_of_units; $i++) {
        Unit::create(['property_id' => $property->id,
                      'unit_number' => "Unit {$i}", 'status' => 'vacant']);
    }
}

// Multi-unit, multi-floor: read JSON floor_config from form
// floor_config = [{"floor":"Ground","count":4}, {"floor":"Floor 1","count":6}]
foreach ($floorConfig as $floor) {
    $letter = chr(65 + $floorIndex);   // A, B, C, ...
    for ($u = 1; $u <= $floor['count']; $u++) {
        Unit::create(['property_id' => $property->id,
                      'unit_number'  => "{$letter}{$u}",    // A1, A2, B1, ...
                      'floor_number' => $floor['floor'],
                      'status'       => 'vacant']);
    }
}""")

doc.add_paragraph()
subheading(doc,"11.2  LandlordController — Dashboard Metrics")
code_block(doc,"""public function dashboard()
{
    $landlordId = Auth::id();

    // Sync all lease payments before computing dashboard metrics
    app(LeasePaymentSyncService::class)->syncLandlord($landlordId);

    $properties   = Property::where('landlord_id', $landlordId)
                             ->with('units','leases.payments')->get();
    $activeLeases = Lease::where('landlord_id', $landlordId)
                         ->where('status','active')->get();

    $monthlyRevenue    = $activeLeases->sum('monthly_rent');
    $overdueAmount     = Payment::whereIn('lease_id', $activeLeases->pluck('id'))
                                 ->where('status','overdue')->sum('amount');
    $collectionRate    = ...;  // paid / (paid+overdue) * 100
    $sixMonthChartData = ...;  // monthly sum of payments for last 6 months

    return view('landlord.dashboard', compact(...));
}""")

doc.add_paragraph()
subheading(doc,"11.3  PaymentController — Index (Most Complex Method)")
para(doc,
    "The index() method is the most complex in the codebase. It syncs all payments, "
    "computes 4 stat counters, applies a filter, paginates 25 tenants per page, "
    "then pre-computes per-tenant flags in a single additional query to avoid N+1 "
    "database hits in the Blade template.")
code_block(doc,"""// 1. Sync all payments for this landlord
app(LeasePaymentSyncService::class)->syncLandlord($landlord->id);

// 2. Build base query of all tenants linked to this landlord
$base = User::where('role','tenant')
    ->where(function($q) use ($propertyIds, $landlord) {
        $q->whereHas('leasesAsTenant', fn($q2) => $q2->whereIn('property_id',$propertyIds))
          ->orWhere('landlord_id', $landlord->id);
    });

// 3. Compute stats (total, pending, overdue, upcoming next-7-days)
$stats = ['total'=>$base->count(), 'pending'=>..., 'overdue'=>..., 'upcoming'=>...];

// 4. Apply selected filter (all/pending/overdue/upcoming/paid/previous)
switch ($filter) { ... }

// 5. Paginate + eager-load leases
$tenants = $query->with(['leasesAsTenant'=>...])->paginate(25);

// 6. Load most-relevant payment per tenant in ONE query, group by tenant
$paymentMap = Payment::whereIn('tenant_id', $tenants->pluck('id'))
    ->orderByRaw("FIELD(status,'overdue','pending','paid')")
    ->get()->groupBy('tenant_id');

// 7. Pre-compute flags in PHP (avoids N+1 in Blade)
foreach ($tenants as $tenant) {
    $payment = $paymentMap->get($tenant->id)?->first();
    $rowData[$tenant->id] = [
        'payment'        => $payment,
        'can_generate'   => $sync->canGenerateControlNumber($payment),
        'active_control' => $sync->activeControlNumber($payment),
    ];
}""")

doc.add_paragraph()
subheading(doc,"11.4  AdminController — Settings & Audit Logs")
code_block(doc,"""// Settings update — handles unchecked booleans (not sent by browser forms)
public function settingsUpdate(Request $request)
{
    $booleanKeys = ['maintenance_mode', 'allow_tenant_registration'];

    foreach ($request->except(['_token','_method']) as $key => $value) {
        $val = in_array($key, $booleanKeys)
             ? ($request->boolean($key) ? '1' : '0')
             : $value;

        SystemSetting::set($key, $val);   // updateOrCreate on system_settings
    }

    AuditLog::log('settings_updated', 'Admin updated system settings');
    return back()->with('success', 'Settings saved.');
}

// Audit logs — paginated with filters
public function auditLogsIndex(Request $request)
{
    $query = AuditLog::with('user')->latest();

    if ($request->filled('action'))  $query->where('action',  $request->action);
    if ($request->filled('user_id')) $query->where('user_id', $request->user_id);
    if ($request->filled('from'))    $query->whereDate('created_at','>=',$request->from);
    if ($request->filled('to'))      $query->whereDate('created_at','<=',$request->to);

    $logs = $query->paginate(50)->withQueryString();
    return view('admin.audit-logs.index', compact('logs'));
}""")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 12. SERVICE LAYER
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"12. Service Layer")
para(doc,
    "Three service classes in app/Services/ encapsulate domain logic that is too "
    "complex for controllers and spans multiple models or external systems.")

subheading(doc,"12.1  LeasePaymentSyncService")
para(doc,
    "This service is the engine of the payment system. It is called at the start "
    "of every payment-related page load to ensure Payment records are up-to-date "
    "with the current date and lease status.")

code_block(doc,"""// syncLandlord — entry point; processes 100 leases at a time
public function syncLandlord(int $landlordId): void
{
    Lease::where('landlord_id', $landlordId)
        ->with('payments')
        ->chunkById(100, function (Collection $leases) {
            $leases->each(fn (Lease $lease) => $this->syncLease($lease));
        });
}

// syncLease — creates or updates a single Payment for the current billing cycle
public function syncLease(Lease $lease): ?Payment
{
    $this->syncPaymentStatuses($lease->payments);

    if (! $this->leaseCanReceivePayment($lease)) {
        return null;   // no tenant, or lease inactive/expired
    }

    $dueDate = $this->currentDueDate($lease);
    $status  = $dueDate->isPast() && ! $dueDate->isToday() ? 'overdue' : 'pending';

    $payment = $lease->payments->whereIn('status',['pending','overdue'])->first();

    if (! $payment) {
        return Payment::create([
            'lease_id'  => $lease->id,
            'tenant_id' => $lease->tenant_id,
            'amount'    => $lease->monthly_rent,
            'due_date'  => $dueDate->toDateString(),
            'status'    => $status,
            'notes'     => 'Generated from lease rent terms.',
        ]);
    }

    // Update but DON'T move due_date if a control number is already active
    $updates = ['tenant_id'=>$lease->tenant_id, 'amount'=>$lease->monthly_rent, 'status'=>$status];
    if (empty($payment->control_number)) {
        $updates['due_date'] = $dueDate->toDateString();
    }
    if ($this->hasChanged($payment, $updates)) {
        $payment->update($updates);
    }
    return $payment->refresh();
}

// currentDueDate — calculates the correct due date for the current month
private function currentDueDate(Lease $lease): Carbon
{
    $today   = now()->startOfDay();
    $day     = $lease->payment_day ?: $lease->start_date->day;
    $dueDate = $today->copy()->day(min($day, $today->daysInMonth));

    if ($dueDate->lt($lease->start_date)) $dueDate = $lease->start_date->copy();
    if ($dueDate->gt($lease->end_date))   $dueDate = $lease->end_date->copy();

    return $dueDate;
}""")

doc.add_paragraph()
subheading(doc,"12.2  NmbPaymentService")
para(doc,
    "Handles all communication with NMB Bank's Zalongwa SPG (Service Payment Gateway). "
    "Bearer tokens are cached for 14 minutes. The postWithAuth() helper automatically "
    "detects a 401 response, clears the cache, and retries once with a fresh token.")
code_block(doc,"""// Token management with 14-minute cache
public function getToken(bool $forceRefresh = false): array
{
    if (! $forceRefresh && Cache::has('nmb_spg_token')) {
        return ['token' => Cache::get('nmb_spg_token')];
    }

    $response = Http::post($this->baseUrl . '/api/v1/login', [
        'client_usr' => $this->clientUsr,
        'client_key' => $this->clientKey,
    ]);

    $token = $response->json('token');
    Cache::put('nmb_spg_token', $token, now()->addMinutes(14));
    return ['token' => $token];
}

// Control number payload
$payload = [
    'systemName'  => 'REMIS',
    'systemCode'  => 'SP1001',
    'payerID'     => 'TENANT-' . $tenant->id,
    'firstName'   => $firstName,
    'lastName'    => $lastName,
    'email'       => $tenant->email,
    'payerMobile' => $this->normalizeMobile($tenant->phone),
    'currency'    => 'TZS',
    'amount'      => (float) $payment->amount,
    'amountType'  => 'EXACT',
    'paymentType' => 'RENT',
    'paymentDesc' => 'RENT payment - PropertyName / UnitNumber',
];

// POST /api/v1/generatectlno  =>  { "status":"success", "reference_number":"1234567890" }""")

insert_image(doc, f"{IMGS}/payment_flow.png", width=6.2,
             caption="Figure 6 – NMB Rent Payment Flow (end-to-end)")

doc.add_paragraph()
subheading(doc,"12.3  BriqSmsService")
para(doc,
    "Sends SMS messages via BRIQ's REST API (Tanzania). The service normalizes "
    "all input phone numbers to E.164 format, accepting four common Tanzanian formats.")
code_block(doc,"""// Accepted formats all normalised to +255XXXXXXXXX:
//   07XXXXXXXX   (10 digits, leading zero)
//   7XXXXXXXX    (9 digits, no prefix)
//   2557XXXXXXXX (12 digits, country code only)
//  +2557XXXXXXXX (full E.164)

private function normalizePhone(string $phone): array
{
    $digits = preg_replace('/\\D/', '', $phone);

    if (strlen($digits) === 9)
        $subscriber = $digits;
    elseif (strlen($digits) === 10 && str_starts_with($digits,'0'))
        $subscriber = substr($digits, 1);
    elseif (strlen($digits) === 12 && str_starts_with($digits,'255'))
        $subscriber = substr($digits, 3);
    else
        return ['valid'=>false, 'error'=>'Invalid Tanzanian mobile number.'];

    // Must start with 6 or 7 (Tanzanian mobile prefixes)
    if (!preg_match('/^[67]\\d{8}$/', $subscriber))
        return ['valid'=>false, 'error'=>'Not a valid Tanzanian mobile number.'];

    return ['valid'=>true, 'number'=>'+255'.$subscriber];
}

// SMS payload sent to BRIQ API
Http::withHeaders(['X-API-Key' => $this->apiKey])->post($endpoint, [
    'content'    => $message,
    'recipients' => [$normalizedNumber],  // without the '+' prefix
    'sender_id'  => 'REMIS',
]);""")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 13. MAIL SYSTEM
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"13. Mail System")
para(doc,
    "Transactional emails are sent via Laravel's Mail facade using dedicated Mailable "
    "classes in app/Mail/. Each class constructs an HTML email from a Blade template "
    "in resources/views/mail/. The development environment uses the log driver "
    "(emails written to storage/logs/laravel.log); production uses SMTP via Gmail.")

add_table(doc,
    ["Mailable Class","Subject","Triggered By","Template"],
    [
        ["TenantInvitation",       "Your REMIS Account is Ready",               "Landlord clicks Invite Tenant",       "mail/tenant-invitation.blade.php"],
        ["PaymentControlNumber",   "Your Rent Payment Control Number",           "Landlord sends control number (email)","mail/payment-control-number.blade.php"],
        ["LeaseTerminationNotice", "Your Lease Has Been Terminated",            "Landlord sends termination notice",   "mail/lease-termination.blade.php"],
        ["TenantWarningNotice",    "Official Warning Notice Regarding Your Lease","Landlord sends warning notice",     "mail/tenant-warning-notice.blade.php"],
    ],
    col_widths=[1.9,2.2,1.8,1.7]
)

doc.add_paragraph()
subheading(doc,"13.1  Tenant Invitation Flow")
para(doc,
    "When a landlord invites a tenant, a signed URL is generated using Laravel's "
    "password broker. The tenant receives an email with a 'Set Password' button "
    "that links to /create-password/{token}?email=..., valid for 7 days (10,080 minutes).")
code_block(doc,"""// LandlordController::tenantInvite
$token = app('auth.password.broker')->createToken($tenant);
$setupUrl = route('password.setup.show', ['token' => $token, 'email' => $tenant->email]);

Mail::to($tenant->email)->send(
    new TenantInvitation($tenant, $plainPasswordHint, $setupUrl)
);

$tenant->update(['invitation_status' => 'invited',
                 'force_password_change' => true]);

// PasswordSetupController::store (tenant clicks link)
$user = User::where('email', $request->email)->firstOrFail();
app('auth.password.broker')->reset([...], function($user, $password) {
    $user->update(['password' => Hash::make($password),
                   'force_password_change' => false]);
});""")

doc.add_paragraph()
subheading(doc,"13.2  Mail Configuration (.env)")
code_block(doc,"""# Development
MAIL_MAILER=log                  # writes to storage/logs/laravel.log

# Production
MAIL_MAILER=smtp
MAIL_HOST=smtp.gmail.com
MAIL_PORT=587
MAIL_ENCRYPTION=tls
MAIL_USERNAME=your@gmail.com
MAIL_PASSWORD=app-specific-password
MAIL_FROM_ADDRESS=noreply@remis.tz
MAIL_FROM_NAME="REMIS Rental System" """)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 14. EXTERNAL API INTEGRATIONS
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"14. External API Integrations")
para(doc,
    "REMIS integrates with three Tanzanian external services: NMB Bank's payment "
    "gateway, BRIQ's SMS provider, and the Tanzania Revenue Authority's TIN lookup API.")

subheading(doc,"14.1  NMB Bank SPG (Zalongwa Payment Gateway)")
add_table(doc,
    ["API Endpoint","Method","Purpose"],
    [
        ["/api/v1/login",             "POST", "Authenticate: obtain 15-min Bearer token (cached 14 min)"],
        ["/api/v1/generatectlno",     "POST", "Generate a control number for a payment (live)"],
        ["/api/v1/generatedemoctlno", "POST", "Generate a demo control number (testing/staging)"],
        ["/api/v1/getpayment",        "POST", "Poll payment status: returns paid=true when tenant has paid"],
        ["/api/v1/getdemopayment",    "POST", "Poll demo payment status"],
        ["/api/v1/verification",      "POST", "Verify a control number is valid and active"],
    ],
    col_widths=[2.2,0.7,3.7]
)
code_block(doc,"""# Required environment variables
NMB_SPG_URL=https://nmb.spg.co.tz
NMB_SPG_USER=client_email@example.com
NMB_SPG_KEY=client_password
NMB_SYSTEM_NAME=REMIS
NMB_SYSTEM_CODE=SP1001
NMB_SPG_TIMEOUT=20
NMB_CA_BUNDLE=/path/to/ca-bundle.crt    # optional custom SSL cert""")

doc.add_paragraph()
para(doc,
    "The service implements automatic token refresh on 401 responses. If the cached "
    "token is rejected, Cache::forget() is called and the request retried once with "
    "a freshly obtained token before returning an error to the caller.",
    size=10)

doc.add_paragraph()
subheading(doc,"14.2  BRIQ SMS (Tanzania Bulk SMS)")
code_block(doc,"""BRIQ_API_KEY=your_api_key
BRIQ_BASE_URL=https://karibu.briq.tz
BRIQ_SENDER_ID=REMIS
BRIQ_BEARER_TOKEN=          # optional secondary auth""")
para(doc,
    "Used to deliver rent payment control numbers to tenants via SMS. The service "
    "normalizes Tanzanian mobile numbers, validates them (must start with 06x or 07x), "
    "and never logs the API key — only sanitized recipient numbers are written to logs.",
    size=10)

doc.add_paragraph()
subheading(doc,"14.3  Tanzania Revenue Authority — TIN Verification")
code_block(doc,"""TRA_API_BASE_URL=https://api.tra.go.tz
TRA_API_KEY=your_api_key
TRA_API_TIMEOUT=15
TRA_DEV_MODE=true      # returns realistic mock data in development""")
para(doc,
    "The TinVerificationController provides an AJAX endpoint that the tenant creation "
    "form calls when a TIN is entered. In dev_mode=true, realistic mock data is returned "
    "instantly without hitting the live TRA API. In production, it queries the real TRA "
    "endpoint and returns normalized fields (taxpayer_id, taxpayer_name, has_vat, vrn).",
    size=10)
code_block(doc,"""// AJAX request from tenant form
POST /landlord/tenants/verify-tin   { tin: "100-123-456" }

// Response (production)
{
    "success":       true,
    "taxpayer_id":   "100-123-456",
    "taxpayer_name": "JOHN DOE LIMITED",
    "has_vat":       false,
    "vrn":           null
}""")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 15. PDF REPORT GENERATION
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"15. PDF Report Generation")
para(doc,
    "PDF documents are generated using the barryvdh/laravel-dompdf package v3.1, "
    "which renders Blade templates to A4 PDF files. Two types of PDFs are produced: "
    "individual lease documents and batch management reports.")

add_table(doc,
    ["Report / Document","Route","Contents"],
    [
        ["Lease PDF",              "GET /landlord/leases/{id}/download",         "Tenant/landlord details, property, dates, rent, security deposit, terms, signatures section"],
        ["Rent Payments Report",   "GET /landlord/reports/pdf/rent-payments",    "All payment records with status, due date, paid date, amount"],
        ["Tenant Inventory",       "GET /landlord/reports/pdf/tenants",          "All tenants with contact info, TIN, NIDA, current lease status"],
        ["Overdue Payments",       "GET /landlord/reports/pdf/overdue",          "Tenants with overdue balances, amounts, days overdue"],
        ["Property Inventory",     "GET /landlord/reports/pdf/properties",       "All properties with unit counts, occupancy rate, monthly revenue"],
    ],
    col_widths=[1.7,2.5,2.4]
)
code_block(doc,"""// Pattern used in all report methods
public function reportRentPaymentsPdf()
{
    $payments = Payment::whereHas('lease',
                    fn($q) => $q->where('landlord_id', Auth::id()))
                ->with(['tenant', 'lease.property'])
                ->orderBy('due_date','desc')
                ->get();

    $pdf = Pdf::loadView('landlord.reports.pdf.rent-payments',
                         compact('payments'));
    $pdf->setPaper('A4', 'portrait');
    return $pdf->stream('rent-payments-' . now()->format('Y-m-d') . '.pdf');
}""")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 16. ARTISAN CONSOLE COMMANDS
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"16. Artisan Console Commands")
para(doc,
    "Two custom Artisan commands in app/Console/Commands/ support server-side "
    "PDF generation outside of the HTTP request cycle.")

add_table(doc,
    ["Command Signature","Purpose","Usage Example"],
    [
        ["remis:generate-pdf {leaseId} {outputPath}",
         "Generate a single lease as a PDF file to disk",
         "php artisan remis:generate-pdf 42 /tmp/lease-42.pdf"],
        ["remis:generate-report-pdf {type} {landlordId} {outputPath}",
         "Generate a batch management report PDF (types: rent-payments, tenants, overdue, properties)",
         "php artisan remis:generate-report-pdf overdue 7 /tmp/overdue.pdf"],
    ],
    col_widths=[2.8,1.9,1.9]
)
para(doc,
    "These commands use GD image processing to embed property images in PDFs. "
    "They are useful for scheduled report generation (e.g., monthly rent summaries "
    "dispatched via email as attachments) without needing an HTTP session.",
    size=10)

# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Part 2 saved: {OUT}")
