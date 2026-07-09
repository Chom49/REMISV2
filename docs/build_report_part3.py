"""
REMIS Backend Report – Part 3  (Sections 17-22 + Conclusion)
Appends to the existing REMIS_Backend_Report.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"c:\xampp\htdocs\RemisV2"
IMGS = os.path.join(BASE, "docs", "images")
PUB  = os.path.join(BASE, "public", "images")
OUT  = os.path.join(BASE, "docs", "REMIS_Backend_Report.docx")

doc = Document(OUT)

def set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),hex_color); tcPr.append(shd)

def add_page_break(doc): doc.add_page_break()

def heading(doc,text,level=1,color="1e3a5f"):
    p=doc.add_heading(text,level=level)
    for r in p.runs: r.font.color.rgb=RGBColor.from_string(color)
    return p

def subheading(doc,text,color="2e6da4"): return heading(doc,text,2,color)
def sub3(doc,text,color="17a2b8"):       return heading(doc,text,3,color)

def para(doc,text="",bold=False,italic=False,size=11,color=None,align=None):
    p=doc.add_paragraph()
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    return p

def code_block(doc,text,size=8):
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"F3F4F6"); pPr.append(shd)
    ind=OxmlElement("w:ind"); ind.set(qn("w:left"),"360")
    ind.set(qn("w:right"),"360"); pPr.append(ind)
    r=p.add_run(text); r.font.name="Courier New"; r.font.size=Pt(size)
    r.font.color.rgb=RGBColor(0x1e,0x3a,0x5f)
    return p

def bullet(doc,text,level=0):
    p=doc.add_paragraph(text,style="List Bullet")
    p.paragraph_format.left_indent=Inches(0.25*(level+1))
    return p

def insert_image(doc,path,width=6.0,caption=None):
    if not os.path.exists(path):
        para(doc,f"[Image not found: {path}]",italic=True,color="999999"); return
    doc.add_picture(path,width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size=Pt(9); r.font.italic=True
            r.font.color.rgb=RGBColor(0x66,0x66,0x66)

def add_table(doc,headers,rows,col_widths=None):
    tbl=doc.add_table(rows=1,cols=len(headers))
    tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=tbl.rows[0].cells
    for i,h in enumerate(headers):
        hc[i].text=h; set_cell_bg(hc[i],"1e3a5f")
        r=hc[i].paragraphs[0].runs[0]
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.bold=True; r.font.size=Pt(9)
        hc[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,row in enumerate(rows):
        cells=tbl.add_row().cells
        bg="F8F9FF" if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(row):
            cells[ci].text=str(val); set_cell_bg(cells[ci],bg)
            cells[ci].paragraphs[0].runs[0].font.size=Pt(9)
    if col_widths:
        for row in tbl.rows:
            for ci,w in enumerate(col_widths): row.cells[ci].width=Inches(w)
    return tbl

# ════════════════════════════════════════════════════════════════════════════
# 17. CONFIGURATION FILES
# ════════════════════════════════════════════════════════════════════════════
add_page_break(doc)
heading(doc,"17. Configuration Files")
para(doc,
    "Laravel configuration files in config/ map environment variables from .env "
    "to structured PHP arrays consumed by the application. All sensitive credentials "
    "remain in .env and are never committed to version control.")

subheading(doc,"17.1  config/auth.php")
code_block(doc,"""'guards' => [
    'web' => ['driver' => 'session', 'provider' => 'users'],
],
'providers' => [
    'users' => ['driver' => 'eloquent', 'model' => App\\Models\\User::class],
],
'passwords' => [
    'users' => [
        'provider' => 'users',
        'table'    => 'password_reset_tokens',
        'expire'   => 10080,   // 7 days — tenant invitation tokens
        'throttle' => 60,      // 60-second resend throttle
    ],
],""")

doc.add_paragraph()
subheading(doc,"17.2  config/services.php  (Third-party API credentials)")
code_block(doc,"""return [
    'tra' => [
        'base_url' => env('TRA_API_BASE_URL'),
        'api_key'  => env('TRA_API_KEY'),
        'timeout'  => env('TRA_API_TIMEOUT', 15),
        'dev_mode' => env('TRA_DEV_MODE', true),
    ],
    'briq' => [
        'api_key'      => env('BRIQ_API_KEY'),
        'base_url'     => env('BRIQ_BASE_URL', 'https://karibu.briq.tz'),
        'sender_id'    => env('BRIQ_SENDER_ID', 'REMIS'),
        'bearer_token' => env('BRIQ_BEARER_TOKEN'),
    ],
    'nmb' => [
        'base_url'    => env('NMB_SPG_URL', 'https://nmb.spg.co.tz'),
        'client_usr'  => env('NMB_SPG_USER'),
        'client_key'  => env('NMB_SPG_KEY'),
        'system_name' => env('NMB_SYSTEM_NAME', 'REMIS'),
        'system_code' => env('NMB_SYSTEM_CODE', 'SP1001'),
        'timeout'     => env('NMB_SPG_TIMEOUT', 20),
        'ca_bundle'   => env('NMB_CA_BUNDLE'),
    ],
];""")

doc.add_paragraph()
subheading(doc,"17.3  .env.example  (Key Variables)")
add_table(doc,
    ["Variable","Default/Example","Description"],
    [
        ["APP_NAME",             "REMIS",                     "Application name shown in emails and UI"],
        ["APP_ENV",              "local",                     "Environment: local / staging / production"],
        ["APP_DEBUG",            "true",                      "Set false in production to hide stack traces"],
        ["DB_DATABASE",          "remis_db",                  "MySQL database name"],
        ["SESSION_DRIVER",       "file",                      "Session storage: file / database / redis"],
        ["MAIL_MAILER",          "log",                       "log (dev) or smtp (production)"],
        ["TRA_DEV_MODE",         "true",                      "false in production to use live TRA API"],
        ["NMB_SPG_URL",          "https://nmb.spg.co.tz",     "NMB bank payment gateway base URL"],
        ["NMB_SYSTEM_CODE",      "SP1001",                    "Merchant system code assigned by NMB"],
        ["BRIQ_SENDER_ID",       "REMIS",                     "SMS sender ID shown on tenant's phone"],
    ],
    col_widths=[1.9,1.8,2.9]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 18. TESTING
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"18. Testing")
para(doc,
    "The test suite uses PHPUnit 11 and covers the two most critical backend components: "
    "the SMS service (which must correctly normalize Tanzanian phone numbers and handle "
    "API errors) and the lease payment sync service (which must compute correct payment "
    "statuses and control number validity).")

add_table(doc,
    ["Test File","Type","Test Cases","Coverage Area"],
    [
        ["tests/Feature/BriqSmsServiceTest.php",          "Feature", "9",  "SMS API, phone normalization, error handling, payload structure"],
        ["tests/Unit/LeasePaymentSyncServiceTest.php",    "Unit",    "8",  "Payment status logic, control number expiry, lease eligibility"],
    ],
    col_widths=[3.0,0.7,0.6,2.3]
)

doc.add_paragraph()
subheading(doc,"18.1  BriqSmsServiceTest — Test Cases")
add_table(doc,
    ["Test Case","Assertion"],
    [
        ["it requires api key to be configured",        "Returns error when BRIQ_API_KEY is empty"],
        ["it normalizes 07XXXXXXXX format",             "0712345678 -> +255712345678"],
        ["it normalizes 9-digit format",                "712345678  -> +255712345678"],
        ["it normalizes 255XXXXXXXXX format",           "255712345678 -> +255712345678"],
        ["it rejects numbers starting with 08",         "Returns valid=false, error message"],
        ["it rejects non-Tanzanian numbers",            "Returns valid=false for invalid lengths"],
        ["it sends correct payload to BRIQ endpoint",   "Asserts content, recipients[], sender_id fields"],
        ["it returns success on HTTP 200",              "Returns ['success' => true, 'data' => ...]"],
        ["it returns error on HTTP 4xx",                "Returns ['success' => false, 'error' => ...]"],
    ],
    col_widths=[2.8,3.8]
)

doc.add_paragraph()
subheading(doc,"18.2  LeasePaymentSyncServiceTest — Test Cases")
add_table(doc,
    ["Test Case","Assertion"],
    [
        ["pending stays pending if due date is today",   "status remains 'pending' when due_date == today"],
        ["payment becomes overdue when due date passes", "status flips to 'overdue' when due_date < today"],
        ["paid payment is never changed",                "status always 'paid' regardless of due_date/lease"],
        ["payment on inactive lease becomes overdue",    "lease.status != 'active' forces 'overdue'"],
        ["control number is active within 7 days",       "activeControlNumber() true if generated < 7 days ago"],
        ["control number expires after 7 days",          "activeControlNumber() false if generated > 7 days ago"],
        ["cannot generate when active control exists",   "canGenerateControlNumber() false if activeControlNumber()"],
        ["can generate when previous has expired",       "canGenerateControlNumber() true after 7-day expiry"],
    ],
    col_widths=[2.9,3.7]
)

doc.add_paragraph()
subheading(doc,"18.3  Running Tests")
code_block(doc,"""# Run all tests
php artisan test

# Run with coverage (requires Xdebug or PCOV)
php artisan test --coverage

# Run only SMS tests
php artisan test --filter BriqSmsService

# Run only sync service tests
php artisan test --filter LeasePaymentSync

# Using Composer script (clears config cache first)
composer test""")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 19. SECURITY IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"19. Security Implementation")
para(doc,
    "Security is implemented at multiple layers: input validation, authentication "
    "hardening, route authorization, object-level authorization, and secure "
    "handling of API credentials.")

subheading(doc,"19.1  Authentication Security")
add_table(doc,
    ["Threat","Mitigation","Implementation"],
    [
        ["Brute-force login",       "Rate limiting",                  "RateLimiter::tooManyAttempts() — 5 unlock attempts, 60s decay"],
        ["Session fixation",        "Session regeneration",           "session()->regenerate() called on every successful login"],
        ["CSRF attacks",            "CSRF token middleware",          "VerifyCsrfToken on all POST/PUT/DELETE routes via web middleware group"],
        ["Password exposure",       "Bcrypt hashing",                 "'password' => 'hashed' cast auto-hashes on assignment"],
        ["Stale browser sessions",  "Lock screen",                    "CheckSessionLock redirects viaRemember() sessions to /lock"],
        ["Weak passwords",          "Strong password policy",         "Password::min(8)->mixedCase()->numbers()->symbols() at registration"],
        ["Unauthorized cross-role access","Role middleware",          "abort(403) in EnsureRole if role doesn't match route group"],
    ],
    col_widths=[1.7,1.5,3.4]
)

doc.add_paragraph()
subheading(doc,"19.2  Object-Level Authorization")
para(doc,
    "Beyond role checks, the application verifies ownership of individual records. "
    "PaymentController::authorizePayment() confirms the landlord owns the property "
    "linked to the payment. TenantController scopes all queries to Auth::user()->id. "
    "LandlordController uses where('landlord_id', Auth::id()) on every query.")

doc.add_paragraph()
subheading(doc,"19.3  API Credential Security")
add_table(doc,
    ["API","Security Measures"],
    [
        ["NMB SPG",   "Credentials in .env only; token cached server-side (never sent to browser); only error status codes logged (no credentials)"],
        ["BRIQ SMS",  "API key in .env only; only sanitized recipient numbers (no keys) written to logs"],
        ["TRA API",   "Bearer token in .env only; dev_mode=true in development prevents live API calls"],
        ["General",   "SSL/TLS enforced on all outbound HTTP calls; custom CA bundle path supported via NMB_CA_BUNDLE"],
    ],
    col_widths=[1.0,5.6]
)

doc.add_paragraph()
subheading(doc,"19.4  Input Validation")
para(doc,
    "All user input is validated at the controller level using Laravel's Validator. "
    "Key rules in use:")
add_table(doc,
    ["Field","Validation Rules"],
    [
        ["name",     "required | string | max:100 | regex:letters/spaces/hyphens/apostrophes only"],
        ["email",    "required | email:rfc | max:150 | unique:users,email"],
        ["phone",    "nullable | regex:/+?[0-9 -]{7,20}/"],
        ["password", "required | confirmed | min:8 | mixedCase | numbers | symbols"],
        ["role",     "required | in:landlord,tenant  (registration only)"],
        ["tin",      "AJAX-validated against TRA API before form submission"],
    ],
    col_widths=[1.0,5.6]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 20. PERFORMANCE CONSIDERATIONS
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"20. Performance Considerations")

subheading(doc,"20.1  Database Indexes")
para(doc,
    "Performance indexes were added in migration 2026_05_12_083721. "
    "These cover the most frequent query patterns in the application.")
add_table(doc,
    ["Table","Index Columns","Query Pattern Served"],
    [
        ["payments",   "(lease_id, status)",     "Filter payments for a lease by status (pending/overdue)"],
        ["payments",   "(tenant_id, status)",    "Load tenant's payment history by status"],
        ["payments",   "(due_date, status)",     "Upcoming/overdue date range queries on payments dashboard"],
        ["leases",     "(landlord_id, status)",  "Landlord dashboard: all active leases"],
        ["leases",     "(property_id, status)",  "Property detail: find active lease for a unit"],
        ["leases",     "(tenant_id, status)",    "Tenant portal: find active lease for the logged-in tenant"],
        ["audit_logs", "(action, created_at)",   "Admin audit log filtering by action type and date range"],
        ["audit_logs", "(model_type, model_id)", "Look up all audit events for a specific record"],
    ],
    col_widths=[1.2,1.8,3.6]
)

doc.add_paragraph()
subheading(doc,"20.2  N+1 Query Prevention")
para(doc,
    "Eloquent's eager loading (->with([...])) is used throughout all controllers. "
    "The most complex case is the Payment index — after paginating 25 tenants, "
    "a single additional query loads all relevant payments grouped by tenant_id "
    "in PHP, avoiding per-row database calls in the Blade template.")
code_block(doc,"""// All 25 tenants' payments loaded in ONE query
$paymentMap = Payment::whereIn('tenant_id', $tenants->pluck('id'))
    ->with(['lease.property', 'lease.unit'])
    ->orderByRaw("FIELD(status,'overdue','pending','paid')")
    ->get()
    ->groupBy('tenant_id');    // PHP grouping, no extra DB round-trips

// Flags computed in PHP loop, not in Blade (avoids N+1)
foreach ($tenants as $tenant) {
    $payment = $paymentMap->get($tenant->id)?->first();
    $rowData[$tenant->id] = [
        'can_generate'   => $sync->canGenerateControlNumber($payment),
        'active_control' => $sync->activeControlNumber($payment),
    ];
}""")

doc.add_paragraph()
subheading(doc,"20.3  Chunked Batch Processing")
para(doc,
    "LeasePaymentSyncService::syncLandlord() uses chunkById(100) to process leases "
    "in batches of 100, preventing memory exhaustion for landlords with large portfolios. "
    "Each chunk is processed and then garbage collected before the next is loaded.")

doc.add_paragraph()
subheading(doc,"20.4  API Token Caching")
para(doc,
    "NMB bearer tokens are cached for 14 minutes (tokens expire at 15 minutes) "
    "using Laravel's Cache facade. The dashboard page for a landlord with 50 tenants "
    "triggers zero NMB API calls for token acquisition on warm cache hits.")

doc.add_paragraph()
subheading(doc,"20.5  Pagination")
para(doc,
    "The Payments index paginates at 25 tenants per page using Laravel's paginate(25) "
    "with withQueryString() to preserve filter parameters across pages. "
    "Audit logs are paginated at 50 records per page.")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 21. DEPLOYMENT NOTES
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"21. Deployment Notes")

subheading(doc,"21.1  Local Development (XAMPP)")
add_table(doc,
    ["Requirement","Specification"],
    [
        ["XAMPP",           "Apache 2.4 + MySQL 8 + PHP 8.2 (Windows)"],
        ["PHP Extensions",  "pdo_mysql, mbstring, openssl, fileinfo, gd, curl, zip (all included in XAMPP)"],
        ["mod_rewrite",     "Must be enabled in httpd.conf; AllowOverride All on htdocs/RemisV2"],
        ["Document Root",   "Set VirtualHost to c:/xampp/htdocs/RemisV2/public"],
        ["Permissions",     "storage/ and bootstrap/cache/ must be writable by Apache"],
    ],
    col_widths=[1.7,4.9]
)

doc.add_paragraph()
subheading(doc,"21.2  Setup Script")
code_block(doc,"""# One-command setup (runs composer install + migrate + npm build)
composer run setup

# Equivalent to:
composer install
php artisan key:generate
php artisan migrate --force
npm install --ignore-scripts
npm run build""")

doc.add_paragraph()
subheading(doc,"21.3  Development Server")
code_block(doc,"""# Start all services concurrently (PHP server + queue + Vite HMR)
composer run dev

# Equivalent to running these in parallel:
php artisan serve              # PHP dev server on :8000
php artisan queue:listen       # Queue worker
npm run dev                    # Vite HMR for frontend assets""")

doc.add_paragraph()
subheading(doc,"21.4  Production Checklist")
add_table(doc,
    ["Step","Command / Setting"],
    [
        ["Set environment",         "APP_ENV=production  APP_DEBUG=false  in .env"],
        ["Enable live TRA API",     "TRA_DEV_MODE=false"],
        ["Enable live NMB API",     "NMB_SPG_URL=https://nmb.spg.co.tz (production URL)"],
        ["Enable SMTP mail",        "MAIL_MAILER=smtp  MAIL_HOST=smtp.gmail.com"],
        ["Cache configuration",     "php artisan config:cache"],
        ["Cache routes",            "php artisan route:cache"],
        ["Cache views",             "php artisan view:cache"],
        ["Run optimizer",           "php artisan optimize"],
        ["Set storage permissions", "chmod -R 775 storage bootstrap/cache"],
        ["Symlink storage",         "php artisan storage:link"],
    ],
    col_widths=[1.8,4.8]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 22. COMPONENT INTERACTION SUMMARY
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"22. Component Interaction Summary")
para(doc,
    "The diagram below traces the complete flow of the most complex backend operation — "
    "a landlord generating a rent control number and sending it to a tenant via SMS. "
    "This single user action touches every layer of the backend.")

insert_image(doc, f"{IMGS}/payment_flow.png", width=6.2,
             caption="Figure 7 – Complete NMB Payment Flow Across All Backend Layers")

doc.add_paragraph()
subheading(doc,"22.1  Step-by-Step Interaction Trace")
add_table(doc,
    ["Step","Component","Action"],
    [
        ["1","Browser","POST /landlord/payments/{id}/generate  (CSRF token included)"],
        ["2","Route (web.php)","Matches route landlord.payments.generate, applies middleware stack"],
        ["3","Middleware: auth","Verifies session is authenticated; redirects to / if not"],
        ["4","Middleware: role:landlord","Checks user.role == 'landlord'; abort(403) if not"],
        ["5","Middleware: locked","Checks !viaRemember() and !locked; redirects to /lock if needed"],
        ["6","PaymentController::generateControlNumber","authorizePayment() confirms landlord owns property"],
        ["7","LeasePaymentSyncService::syncPaymentStatuses","Recalculates payment status (may flip pending to overdue)"],
        ["8","LeasePaymentSyncService::canGenerateControlNumber","Validates: not paid, lease active, no active control number"],
        ["9","NmbPaymentService::getToken","Returns cached token or calls NMB /api/v1/login"],
        ["10","NmbPaymentService::generateControlNumber","POST /api/v1/generatectlno to NMB Bank SPG"],
        ["11","NMB Bank SPG","Returns { status:'success', reference_number:'1234567890' }"],
        ["12","Payment::update","Saves control_number and control_number_generated_at to DB"],
        ["13","Browser (redirect)","Flash success: 'Control number generated: 1234567890'"],
        ["14","Landlord clicks Send via SMS","POST /landlord/payments/{id}/send  channel=sms"],
        ["15","PaymentController::sendControlNumber","Validates active control, loads tenant and lease"],
        ["16","BriqSmsService::send","Normalizes tenant phone to E.164, POSTs to BRIQ API"],
        ["17","BRIQ API","Delivers SMS to tenant's Tanzanian mobile number"],
        ["18","Payment::update","Records control_number_sent_at and control_number_sent_via='sms'"],
        ["19","Tenant pays via NMB","Tenant visits NMB branch/ATM/app, enters control number"],
        ["20","PaymentController::pollAll (AJAX)","Calls NMB /api/v1/getpayment for each pending payment"],
        ["21","NMB Bank SPG","Returns paid=true with transaction reference, receipt, payer details"],
        ["22","Payment::update","status='paid', nmb_transaction_id, nmb_receipt_number, nmb_paid_at saved"],
        ["23","Dashboard","Payment shows as confirmed; landlord sees receipt and payer name"],
    ],
    col_widths=[0.4,2.2,4.0]
)

doc.add_paragraph()
insert_image(doc, f"{IMGS}/lease_lifecycle.png", width=6.0,
             caption="Figure 8 – Lease Lifecycle State Machine and Payment Status Flow")

doc.add_paragraph()
subheading(doc,"22.2  Final Architecture Summary")
para(doc,
    "REMIS V2 demonstrates a cohesive, layered backend where each concern is "
    "handled by the right component:")
bullet(doc,"Routes (web.php) — URL-to-controller mapping and middleware grouping")
bullet(doc,"Middleware — cross-cutting concerns (auth, role, session lock, force-password)")
bullet(doc,"Controllers — HTTP request/response handling, input validation, delegation")
bullet(doc,"Service Layer — complex domain logic (payment sync, NMB, SMS, TIN)")
bullet(doc,"Models — data structure, relationships, casts, domain helper methods")
bullet(doc,"Migrations — version-controlled schema evolution across 4 phases")
bullet(doc,"Mail classes — decoupled, testable transactional email objects")
bullet(doc,"Config / .env — environment-specific credentials and feature flags")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"Conclusion")
para(doc,
    "The REMIS V2 backend is a production-ready Laravel 12 application purpose-built "
    "for the Tanzanian rental property market. It was developed in four phases from "
    "January 2024 to June 2026, evolving from a basic property management system to "
    "a comprehensive platform with deep integrations into Tanzania's financial and "
    "regulatory infrastructure.")

doc.add_paragraph()
para(doc,
    "The architecture prioritizes data isolation (each landlord sees only their own "
    "data), security (rate-limiting, session locking, object-level authorization), "
    "and maintainability (service objects for complex logic, named routes for "
    "refactoring safety, and descriptive migration naming for schema traceability).")

doc.add_paragraph()
para(doc,
    "The NMB Bank SPG integration enables a complete digital rent collection workflow: "
    "control number generation, multi-channel delivery (email or SMS), and "
    "real-time payment confirmation polling — all with automatic token management "
    "and graceful error handling.")

doc.add_paragraph()
para(doc,
    "With 28 database migrations, 9 Eloquent models, 7 controllers, 3 service classes, "
    "4 mail classes, 3 middleware, and 60+ named routes, REMIS V2 represents a "
    "complete, well-structured backend implementation that can be maintained and "
    "extended by any Laravel developer familiar with the framework.",
    italic=True)

# ── final screenshots ────────────────────────────────────────────────────────
add_page_break(doc)
heading(doc,"Appendix — Application Screenshots")

insert_image(doc, os.path.join(PUB,"landingPage","remis1.png"), width=5.8,
             caption="Figure A1 – REMIS Landing Page")

doc.add_paragraph()
insert_image(doc, os.path.join(PUB,"signIn","signin1.png"), width=5.8,
             caption="Figure A2 – Sign-In / Lock Screen")

doc.add_paragraph()
insert_image(doc, f"{IMGS}/architecture.png", width=5.8,
             caption="Figure A3 – Complete System Architecture Diagram")

doc.add_paragraph()
insert_image(doc, f"{IMGS}/erd.png", width=5.8,
             caption="Figure A4 – Full Database Entity-Relationship Diagram")

doc.add_paragraph()
insert_image(doc, f"{IMGS}/request_lifecycle.png", width=5.8,
             caption="Figure A5 – HTTP Request Lifecycle & Middleware Pipeline")

# ── save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)

size_mb = os.path.getsize(OUT) / 1024 / 1024
print(f"Final report saved: {OUT}")
print(f"File size: {size_mb:.2f} MB")
