"""
REMIS Backend Report – Word Document Generator
Builds docs/REMIS_Backend_Report.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE   = r"c:\xampp\htdocs\RemisV2"
IMGS   = os.path.join(BASE, "docs", "images")
PUB    = os.path.join(BASE, "public", "images")
OUT    = os.path.join(BASE, "docs", "REMIS_Backend_Report.docx")

# ── helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"),   "single")
        tag.set(qn("w:sz"),    "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color","CCCCCC"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_page_break(doc):
    doc.add_page_break()

def heading(doc, text, level=1, color="1e3a5f"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def subheading(doc, text, color="2e6da4"):
    return heading(doc, text, level=2, color=color)

def sub3(doc, text, color="17a2b8"):
    return heading(doc, text, level=3, color=color)

def para(doc, text="", bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def code_block(doc, text, size=8):
    """Monospace shaded paragraph for code snippets."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F3F4F6")
    pPr.append(shd)
    # indent
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),  "360")
    ind.set(qn("w:right"), "360")
    pPr.append(ind)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1e,0x3a,0x5f)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level+1))
    return p

def numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 * (level+1))
    return p

def insert_image(doc, path, width=6.0, caption=None):
    if not os.path.exists(path):
        para(doc, f"[Image not found: {path}]", italic=True, color="999999")
        return
    doc.add_picture(path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size  = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66,0x66,0x66)

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "1e3a5f")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        run.font.bold = True
        run.font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        bg = "F8F9FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            set_cell_bg(cells[ci], bg)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return tbl

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pb.append(bot)
    pPr.append(pb)

# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── page margins ─────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
logo = os.path.join(PUB,"signIn","logo_transparent.png")
insert_image(doc, logo, width=2.0)

doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("REMIS V2")
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1e,0x3a,0x5f)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Rental Management Information System")
r2.font.size  = Pt(18)
r2.font.color.rgb = RGBColor(0x2e,0x6d,0xa4)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("BACKEND IMPLEMENTATION REPORT")
r3.font.size  = Pt(22)
r3.font.bold  = True
r3.font.color.rgb = RGBColor(0x17,0xa2,0xb8)

doc.add_paragraph()
insert_image(doc, os.path.join(PUB,"landingPage","remis1.png"), width=5.5)

doc.add_paragraph()
for line in [
    ("Framework:",   "Laravel 12  |  PHP 8.2+"),
    ("Database:",    "MySQL 8  (XAMPP / Windows)"),
    ("Report Date:", "June 2026"),
    ("Version:",     "2.0  (commit 2d39813)"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(line[0] + "  ")
    rb.bold = True
    rb.font.size = Pt(11)
    rv = p.add_run(line[1])
    rv.font.size = Pt(11)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"Table of Contents", level=1)
toc_entries = [
    ("1",  "Project Overview"),
    ("2",  "Technology Stack"),
    ("3",  "Backend Architecture"),
    ("4",  "Directory Structure"),
    ("5",  "Database Design"),
    ("6",  "Models and Relationships"),
    ("7",  "Database Migrations"),
    ("8",  "Authentication and Authorization"),
    ("9",  "Middleware"),
    ("10", "Routing"),
    ("11", "Controllers"),
    ("12", "Service Layer"),
    ("13", "Mail System"),
    ("14", "External API Integrations"),
    ("15", "PDF Report Generation"),
    ("16", "Artisan Console Commands"),
    ("17", "Configuration Files"),
    ("18", "Testing"),
    ("19", "Security Implementation"),
    ("20", "Performance Considerations"),
    ("21", "Deployment Notes"),
    ("22", "Component Interaction Summary"),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rb = p.add_run(f"  {num}. ")
    rb.bold = True
    rb.font.size = Pt(10)
    rb.font.color.rgb = RGBColor(0x1e,0x3a,0x5f)
    rt = p.add_run(title)
    rt.font.size = Pt(10)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"1. Project Overview")
para(doc,
    "REMIS V2 (Rental Management Information System) is a full-stack web application "
    "built on Laravel 12 for the Tanzanian rental property market. It provides landlords "
    "with a unified platform to manage properties, units, tenants, leases, and rent "
    "payments — deeply integrated with Tanzanian financial and regulatory infrastructure "
    "including NMB Bank, BRIQ SMS, and the Tanzania Revenue Authority (TRA).")

doc.add_paragraph()
sub3(doc,"Core Capabilities")
add_table(doc,
    ["Domain","Capability"],
    [
        ["Property Management",  "Single-unit and multi-unit (multi-floor) property registration with image upload"],
        ["Tenant Management",    "Tenant onboarding, identity fields (TIN/NIDA/gender/nationality), invitations"],
        ["Lease Lifecycle",      "Create, assign tenant, terminate with reason, renew, download as PDF"],
        ["Rent Payments",        "NMB Bank SPG: generate control numbers, send via email/SMS, poll status"],
        ["Notifications",        "SMTP email (Gmail-ready) + BRIQ bulk SMS (Tanzania)"],
        ["TIN Verification",     "Real-time AJAX lookup against Tanzania Revenue Authority API"],
        ["PDF Reports",          "Rent payments, tenants, overdue payments, property inventory"],
        ["Admin Panel",          "System settings, audit logs, database backups, user management"],
        ["Session Security",     "Lock screen, rate-limited unlock, force-password-change for new tenants"],
        ["Multi-role Access",    "Landlord, Tenant, and Admin — each with isolated dashboard and permissions"],
    ],
    col_widths=[1.8, 4.7]
)

doc.add_paragraph()
sub3(doc,"Development Timeline")
add_table(doc,
    ["Phase","Period","Scope"],
    [
        ["Phase 1 – Foundation",         "January 2024",   "Users, Properties, Leases, Payments, Maintenance, Auth"],
        ["Phase 2 – Multi-Unit Support",  "April–May 2026", "Units table, floor layouts, lease lifecycle, identity fields"],
        ["Phase 3 – Security & Admin",    "May 2026",       "Admin panel, audit logs, backups, session lock, role:admin"],
        ["Phase 4 – Payment & Profiles",  "May–Jun 2026",   "NMB SPG integration, user profiles, preferences JSON"],
    ],
    col_widths=[2.0,1.6,3.0]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"2. Technology Stack")
para(doc,
    "REMIS V2 is built on the Laravel 12 framework using PHP 8.2+. The full dependency "
    "list is managed by Composer and defined in composer.json.")

subheading(doc,"2.1  Production Dependencies")
add_table(doc,
    ["Package","Version","Purpose"],
    [
        ["laravel/framework", "^12.0", "MVC framework — routing, ORM, auth, mail, cache, HTTP client"],
        ["php",               "^8.2",  "Runtime — named args, enums, readonly props, fibers"],
        ["barryvdh/laravel-dompdf","^3.1","PDF generation from Blade templates (A4 lease & report PDFs)"],
        ["laravel/tinker",    "^2.9",  "REPL for debugging and quick model inspection in development"],
    ],
    col_widths=[2.0,1.0,3.5]
)

doc.add_paragraph()
subheading(doc,"2.2  Development Dependencies")
add_table(doc,
    ["Package","Version","Purpose"],
    [
        ["phpunit/phpunit",      "^11.0", "Unit and feature test runner"],
        ["fakerphp/faker",       "^1.23", "Fake data generation for factories/seeders"],
        ["laravel/pint",         "^1.13", "PSR-12 code style enforcement"],
        ["mockery/mockery",      "^1.6",  "Mock objects for unit tests"],
        ["nunomaduro/collision", "^8.0",  "Pretty CLI error reporting during test runs"],
    ],
    col_widths=[2.0,1.0,3.5]
)

doc.add_paragraph()
subheading(doc,"2.3  Infrastructure Components")
add_table(doc,
    ["Layer","Technology","Notes"],
    [
        ["Web Server",   "Apache 2.4 (XAMPP)",   "mod_rewrite enabled; AllowOverride All on document root"],
        ["Database",     "MySQL 8.0",            "InnoDB engine; foreign key constraints ON; charset utf8mb4"],
        ["PHP Runtime",  "PHP 8.2.12",           "Extensions: pdo_mysql, mbstring, openssl, fileinfo, gd, curl, zip"],
        ["Session Store","File-based (default)",  "SESSION_DRIVER=file; configurable to Redis or DB"],
        ["Cache",        "File-based (default)",  "Used by NmbPaymentService for 14-min token caching"],
        ["Mail",         "SMTP / Log driver",     "Development: log driver. Production: Gmail SMTP (port 587/TLS)"],
        ["Frontend Build","Vite + Tailwind CSS",  "Managed separately; Blade templates served from resources/views"],
        ["SMS Gateway",  "BRIQ (Tanzania)",       "REST API at karibu.briq.tz; Tanzanian E.164 number normalization"],
        ["Payment GW",   "NMB Bank SPG (Zalongwa)","REST API at nmb.spg.co.tz; Bearer token auth; 14-min TTL"],
        ["TIN Verify",   "TRA API (Tanzania)",    "Tanzania Revenue Authority; dev-mode mock available"],
    ],
    col_widths=[1.5,1.8,3.2]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3. BACKEND ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"3. Backend Architecture")
para(doc,
    "REMIS follows the standard Laravel MVC architecture augmented with a dedicated "
    "Service Layer for business logic that spans multiple models. The application uses "
    "session-based (stateful) authentication rather than a separate REST API, meaning "
    "all routes — including AJAX endpoints — are protected by the same middleware pipeline.")

doc.add_paragraph()
insert_image(doc, f"{IMGS}/architecture.png", width=6.2,
             caption="Figure 1 – REMIS V2 Backend System Architecture")

doc.add_paragraph()
subheading(doc,"3.1  Design Patterns")
add_table(doc,
    ["Pattern","Where Applied","Benefit"],
    [
        ["MVC",                    "All routes → controllers → views",           "Clear separation of concerns"],
        ["Service Objects",        "LeasePaymentSyncService, NmbPaymentService, BriqSmsService","Complex domain logic isolated from controllers"],
        ["Eloquent Active Record", "All 9 models (User, Property, Unit, Lease…)","Readable query API, relationship graph"],
        ["Mailable Classes",       "4 mail classes (Invitation, ControlNumber, Termination, Warning)","Testable, reusable email objects"],
        ["Middleware Pipeline",    "auth → role:X → locked → force.password.change","Stackable request interceptors"],
        ["Named Routes",           "landlord.payments.index, tenant.dashboard…", "Decoupled URL generation, safe refactoring"],
        ["Repository-like Scoping","All queries scoped to Auth::id() via where('landlord_id'…)","Automatic tenant/landlord data isolation"],
    ],
    col_widths=[1.8,2.8,2.0]
)

doc.add_paragraph()
subheading(doc,"3.2  Request Lifecycle")
insert_image(doc, f"{IMGS}/request_lifecycle.png", width=6.0,
             caption="Figure 2 – HTTP Request Lifecycle and Middleware Pipeline")
para(doc,
    "Every HTTP request passes through Laravel's Kernel, which applies the 'web' "
    "middleware group (sessions, CSRF verification, cookie handling) before reaching "
    "route-specific middleware. Authentication is checked first; if the user is "
    "unauthenticated they are redirected to the landing page. The EnsureRole middleware "
    "then verifies the user's role matches the route group. The CheckSessionLock "
    "middleware fires for remembered sessions, redirecting to the lock screen. Finally, "
    "ForcePasswordChange applies only to tenant routes where the flag is set.")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 4. DIRECTORY STRUCTURE
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"4. Directory Structure")
para(doc,
    "The project follows the standard Laravel directory layout with a few custom "
    "additions. The most important backend-specific directories are described below.")

code_block(doc, """RemisV2/
├── app/
│   ├── Console/Commands/          # CLI: GenerateLeasePdf, GenerateReportPdf
│   ├── Http/
│   │   ├── Controllers/           # 7 controllers (one per domain)
│   │   └── Middleware/            # EnsureRole, ForcePasswordChange, CheckSessionLock
│   ├── Mail/                      # 4 Mailable classes
│   ├── Models/                    # 9 Eloquent models
│   ├── Providers/                 # AppServiceProvider, AuthServiceProvider
│   └── Services/                  # LeasePaymentSyncService, NmbPaymentService, BriqSmsService
├── bootstrap/app.php              # Middleware alias registration
├── config/
│   ├── auth.php                   # Guard / provider / password-reset config
│   ├── database.php               # MySQL connection
│   ├── mail.php                   # SMTP config
│   └── services.php               # NMB / BRIQ / TRA credentials map
├── database/
│   ├── migrations/                # 28 migration files (4 phases)
│   └── seeders/                   # SystemSettings seeder (8 default keys)
├── routes/web.php                 # All 60+ named routes
├── resources/views/               # Blade templates (frontend)
├── tests/
│   ├── Feature/BriqSmsServiceTest.php
│   └── Unit/LeasePaymentSyncServiceTest.php
├── composer.json                  # PHP dependencies
└── .env.example                   # Environment variable template""")

doc.add_paragraph()
add_table(doc,
    ["Directory / File","Purpose"],
    [
        ["app/Http/Controllers/",         "One controller per domain. No fat controllers — complex logic delegated to Services."],
        ["app/Http/Middleware/",           "3 custom middleware: role guard, session lock, force-password-change."],
        ["app/Models/",                   "9 Eloquent models covering all business entities."],
        ["app/Services/",                 "3 service classes encapsulating NMB payment, SMS, and lease sync logic."],
        ["app/Mail/",                     "4 Mailable classes for transactional email (invitation, control number, notices)."],
        ["app/Console/Commands/",         "2 Artisan commands for server-side PDF generation."],
        ["database/migrations/",          "28 migration files evolving the schema from 2024 to 2026."],
        ["routes/web.php",                "All application routes — no api.php (AJAX endpoints use web session auth)."],
        ["config/services.php",           "Central registry for all third-party API credentials."],
    ],
    col_widths=[2.2,4.4]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 5. DATABASE DESIGN
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"5. Database Design")
para(doc,
    "The database uses MySQL 8 with InnoDB engine, enforcing foreign key constraints "
    "for referential integrity. The schema evolved across 28 migrations organized in "
    "4 development phases. The central entities are users, properties, units, leases, "
    "and payments.")

doc.add_paragraph()
insert_image(doc, f"{IMGS}/erd.png", width=6.4,
             caption="Figure 3 – Entity-Relationship Diagram (all tables and foreign keys)")

doc.add_paragraph()
subheading(doc,"5.1  Table Summary")
add_table(doc,
    ["Table","Primary Role","Key Foreign Keys"],
    [
        ["users",                 "Stores all roles in one table (landlord/tenant/admin)",  "landlord_id → users.id (self-ref)"],
        ["properties",            "Physical property records per landlord",                 "landlord_id → users.id"],
        ["units",                 "Individual rentable units within properties",             "property_id → properties.id"],
        ["leases",                "Tenant-unit contracts with full lifecycle tracking",      "property_id, unit_id, tenant_id, landlord_id, renewal_of_id"],
        ["payments",              "Rent payment records including NMB transaction data",    "lease_id, tenant_id"],
        ["maintenance_requests",  "Tenant-submitted maintenance tickets",                   "property_id, tenant_id"],
        ["audit_logs",            "Admin action trail with IP and user agent",              "user_id → users.id"],
        ["system_settings",       "Admin-configurable key-value application settings",      "—"],
        ["backups",               "Database backup metadata records",                       "created_by → users.id"],
        ["sessions",              "Laravel session storage (file or DB)",                   "user_id → users.id"],
        ["password_reset_tokens", "Tenant invitation and password reset tokens",            "email → users.email"],
    ],
    col_widths=[1.8,2.5,2.3]
)

doc.add_paragraph()
subheading(doc,"5.2  Key Design Decisions")

sub3(doc,"Unified Users Table")
para(doc,
    "All three roles (landlord, tenant, admin) are stored in a single users table "
    "differentiated by the 'role' column. This simplifies authentication — one guard, "
    "one provider — while role-specific fields (tin, nida_number for tenants; "
    "profile_picture, preferences for all) are nullable on rows where they don't apply.")

sub3(doc,"Self-Referential landlord_id on Users")
para(doc,
    "When a landlord creates a tenant account, the tenant row records the landlord's "
    "ID in the landlord_id column. This creates a direct ownership link, allowing "
    "queries like user.createdTenants() without joining through the leases table.")

sub3(doc,"Units Table Introduced in Phase 2")
para(doc,
    "The original schema linked leases directly to properties (1:1 for single-unit). "
    "Phase 2 introduced the units table to support multi-unit buildings with individual "
    "floor/unit-number tracking. unit_id on leases is nullable to preserve backward "
    "compatibility with legacy single-unit lease records.")

sub3(doc,"JSON preferences Column")
para(doc,
    "User preferences (notification settings, UI options) are stored in a JSON column "
    "cast to a PHP array by Eloquent. This avoids a separate user_preferences table "
    "and allows extending settings without additional migrations.")

sub3(doc,"NMB Payment Fields on payments Table")
para(doc,
    "Rather than a separate nmb_transactions table, all NMB gateway data "
    "(control_number, nmb_transaction_id, nmb_receipt_number, nmb_payer_name, "
    "nmb_payer_mobile, nmb_paid_at) is stored directly on the payments row. "
    "This keeps the payment record self-contained for reporting.")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 6. MODELS AND RELATIONSHIPS
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"6. Models and Relationships")
para(doc,
    "All business entities are represented as Eloquent models in app/Models/. "
    "Each model defines its fillable fields, casts, relationships, and any "
    "domain helper methods.")

subheading(doc,"6.1  User Model  (app/Models/User.php)")
para(doc,
    "The User model extends Authenticatable, giving it Laravel's built-in auth "
    "capabilities (remember tokens, password reset). It serves all three roles.")
code_block(doc,"""class User extends Authenticatable
{
    protected $fillable = [
        'name', 'email', 'phone', 'role', 'password', 'landlord_id',
        'tenant_status', 'tin', 'nida_number', 'gender', 'nationality',
        'force_password_change', 'default_password_hint', 'invitation_status',
        'profile_picture', 'preferences',
    ];

    protected function casts(): array
    {
        return [
            'password'    => 'hashed',      // auto-bcrypt on assignment
            'preferences' => 'array',        // JSON column ↔ PHP array
        ];
    }

    public function isLandlord(): bool { return $this->role === 'landlord'; }
    public function isTenant(): bool   { return $this->role === 'tenant'; }
    public function isAdmin(): bool    { return $this->role === 'admin'; }

    // Self-referential: landlords own tenants they created
    public function createdTenants()   { return $this->hasMany(User::class, 'landlord_id'); }
    public function createdByLandlord(){ return $this->belongsTo(User::class, 'landlord_id'); }

    public function properties()       { return $this->hasMany(Property::class, 'landlord_id'); }
    public function leasesAsLandlord() { return $this->hasMany(Lease::class, 'landlord_id'); }
    public function leasesAsTenant()   { return $this->hasMany(Lease::class, 'tenant_id'); }
    public function payments()         { return $this->hasMany(Payment::class, 'tenant_id'); }
}""")

doc.add_paragraph()
subheading(doc,"6.2  Property Model  (app/Models/Property.php)")
code_block(doc,"""class Property extends Model
{
    protected $fillable = [
        'landlord_id', 'name', 'address', 'city', 'county', 'total_area',
        'type', 'bedrooms', 'bathrooms', 'status', 'description',
        'property_category',   // 'single' | 'multi'
        'number_of_units', 'floor_layout', 'image',
    ];

    public function units()       { return $this->hasMany(Unit::class); }
    public function leases()      { return $this->hasMany(Lease::class); }
    public function activeLease() { return $this->hasOne(Lease::class)
                                             ->where('status','active'); }

    public function isMultiUnit(): bool  { return $this->property_category === 'multi'; }
    public function occupiedUnitsCount() { return $this->units()->where('status','occupied')->count(); }
    public function vacantUnitsCount()   { return $this->units()->where('status','vacant')->count(); }
}""")

doc.add_paragraph()
subheading(doc,"6.3  Lease Model  (app/Models/Lease.php)")
para(doc,
    "The Lease is the central entity binding a tenant, a unit, and a landlord. "
    "It tracks the full contract lifecycle including renewals via a self-referential "
    "renewal_of_id foreign key.")
code_block(doc,"""class Lease extends Model
{
    protected $fillable = [
        'property_id', 'unit_id', 'tenant_id', 'landlord_id',
        'start_date', 'end_date', 'monthly_rent', 'security_deposit',
        'payment_day', 'payment_frequency', 'status',
        'termination_reason', 'termination_notes', 'terminated_at',
        'renewal_of_id',   // Self-FK — points to the lease being renewed
        'lease_terms',
    ];

    protected function casts(): array
    {
        return ['start_date' => 'date', 'end_date' => 'date'];
    }

    public function renewedFrom() { return $this->belongsTo(Lease::class, 'renewal_of_id'); }
    public function renewals()    { return $this->hasMany(Lease::class,   'renewal_of_id'); }

    public function isActive(): bool     { return $this->status === 'active'; }
    public function isTerminated(): bool { return $this->status === 'terminated'; }

    public function daysUntilExpiry(): int
    {
        return (int) now()->startOfDay()->diffInDays($this->end_date, false);
    }
}""")

doc.add_paragraph()
subheading(doc,"6.4  Payment Model  (app/Models/Payment.php)")
para(doc,
    "Payment records are auto-generated by the LeasePaymentSyncService. NMB transaction "
    "data is stored directly on the payment row when payment is confirmed.")
code_block(doc,"""class Payment extends Model
{
    protected $fillable = [
        'lease_id', 'tenant_id', 'amount', 'due_date', 'paid_date',
        'status', 'reference', 'notes',
        // NMB fields (populated after gateway interaction):
        'control_number', 'control_number_generated_at',
        'control_number_sent_at', 'control_number_sent_via',
        'nmb_transaction_id', 'nmb_receipt_number',
        'nmb_payer_name', 'nmb_payer_mobile', 'nmb_paid_at',
    ];

    protected function casts(): array
    {
        return [
            'due_date'                    => 'date',
            'paid_date'                   => 'date',
            'control_number_generated_at' => 'datetime',
            'nmb_paid_at'                 => 'datetime',
        ];
    }
}""")

doc.add_paragraph()
subheading(doc,"6.5  AuditLog Model  (app/Models/AuditLog.php)")
para(doc,
    "All admin actions are recorded using a static helper method, capturing user, "
    "action type, related model, IP address, and user agent.")
code_block(doc,"""class AuditLog extends Model
{
    protected $fillable = [
        'user_id', 'action', 'model_type', 'model_id',
        'description', 'ip_address', 'user_agent',
    ];

    public static function log(string $action, string $description,
                               ?string $modelType = null, ?int $modelId = null): void
    {
        static::create([
            'user_id'     => Auth::id(),
            'action'      => $action,
            'model_type'  => $modelType,
            'model_id'    => $modelId,
            'description' => $description,
            'ip_address'  => request()->ip(),
            'user_agent'  => request()->userAgent(),
        ]);
    }
}""")

doc.add_paragraph()
subheading(doc,"6.6  Model Relationship Map")
add_table(doc,
    ["Model","Has Many","Belongs To"],
    [
        ["User (landlord)", "Property, Lease (as landlord), User (created tenants)", "—"],
        ["User (tenant)",   "Lease (as tenant), Payment, MaintenanceRequest",        "User (landlord via landlord_id)"],
        ["Property",        "Unit, Lease, MaintenanceRequest",                       "User (landlord)"],
        ["Unit",            "Lease",                                                 "Property"],
        ["Lease",           "Payment, Lease (renewals)",                             "Property, Unit, User×2, Lease (renewedFrom)"],
        ["Payment",         "—",                                                     "Lease, User (tenant)"],
        ["MaintenanceRequest","—",                                                   "Property, User (tenant)"],
        ["AuditLog",        "—",                                                     "User"],
    ],
    col_widths=[1.5,2.8,2.3]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 7. MIGRATIONS
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"7. Database Migrations")
para(doc,
    "Database schema changes are managed via 28 Laravel migration files grouped "
    "into four development phases. Running php artisan migrate applies them "
    "in chronological order, building the full schema from scratch.")

subheading(doc,"Phase 1 — Foundation (January 2024)")
add_table(doc,
    ["Migration","What It Creates / Changes"],
    [
        ["0001_01_01_000000_create_users_table",               "users, password_reset_tokens, sessions, cache tables"],
        ["2024_01_01_000010_create_properties_table",           "properties with landlord_id FK"],
        ["2024_01_01_000011_add_county_area_to_properties",     "Adds county, total_area columns"],
        ["2024_01_01_000020_create_leases_table",               "leases with property/tenant/landlord FKs"],
        ["2024_01_01_000021_add_payment_fields_to_leases",      "payment_day, payment_frequency, lease_expiry_reminder_days"],
        ["2024_01_01_000030_create_payments_table",             "payments with lease/tenant FKs, status, reference"],
        ["2024_01_01_000040_create_maintenance_requests_table", "maintenance_requests with priority/status"],
        ["2024_01_01_000041_update_maintenance_requests_table", "Adds viewable_by column"],
    ],
    col_widths=[3.2,3.4]
)

doc.add_paragraph()
subheading(doc,"Phase 2 — Multi-Unit Support (April–May 2026)")
add_table(doc,
    ["Migration","Change"],
    [
        ["2026_04_25_make_tenant_id_nullable_on_leases",          "tenant_id becomes nullable (lease can exist before tenant assigned)"],
        ["2026_04_27_add_landlord_id_to_users_table",             "Self-referential FK for landlord→tenant ownership"],
        ["2026_05_07_000001_add_property_category_to_properties", "property_category, number_of_units, image columns"],
        ["2026_05_07_000002_create_units_table",                  "New units table (unit_number, floor_number, status, notes)"],
        ["2026_05_07_000003_add_unit_id_and_terms_to_leases",     "unit_id FK on leases, lease_terms text field"],
        ["2026_05_07_100001_add_tenant_status_to_users",          "tenant_status column (active/inactive)"],
        ["2026_05_07_100002_add_lifecycle_fields_to_leases",      "termination_reason, termination_notes, terminated_at, renewal_of_id"],
        ["2026_05_07_200001_add_identity_fields_to_users",        "tin, nida_number, gender, nationality (Tanzania compliance)"],
        ["2026_05_12_add_floor_fields",                           "floor_layout on properties, floor_number on units"],
    ],
    col_widths=[3.2,3.4]
)

doc.add_paragraph()
subheading(doc,"Phase 3 — Security & Admin (May 2026)")
add_table(doc,
    ["Migration","Change"],
    [
        ["2026_05_11_000001_add_force_password_change",       "force_password_change bool, default_password_hint string"],
        ["2026_05_11_000002_add_invitation_status",           "invitation_status on users (pending/invited/accepted)"],
        ["2026_05_12_add_performance_indexes",                "Composite indexes on payments and leases for query optimization"],
        ["2026_05_29_000001_add_admin_role",                  "Extends role enum to include 'admin'"],
        ["2026_05_29_000002_create_audit_logs_table",         "audit_logs with indexed (action, created_at) and (model_type, model_id)"],
        ["2026_05_29_000003_create_system_settings_table",    "system_settings key-value table, pre-seeded with 8 defaults"],
        ["2026_05_29_000004_create_backups_table",            "backups metadata table with created_by FK"],
    ],
    col_widths=[3.2,3.4]
)

doc.add_paragraph()
subheading(doc,"Phase 4 — User Profiles & NMB Payments (May–June 2026)")
add_table(doc,
    ["Migration","Change"],
    [
        ["2026_05_31_add_profile_settings_to_users",          "profile_picture path, preferences JSON column"],
        ["2026_06_02_add_nmb_fields_to_payments_table",       "All NMB gateway fields: control_number, nmb_transaction_id, nmb_receipt_number, nmb_payer_* , nmb_paid_at"],
    ],
    col_widths=[3.2,3.4]
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 8. AUTHENTICATION AND AUTHORIZATION
# ════════════════════════════════════════════════════════════════════════════
heading(doc,"8. Authentication and Authorization")

insert_image(doc, f"{IMGS}/auth_flow.png", width=6.0,
             caption="Figure 4 – Role-Based Authentication and Authorization Flow")

subheading(doc,"8.1  Login Flow  (AuthController::login)")
para(doc,
    "Authentication uses Laravel's Auth::attempt() with remember-me always enabled. "
    "On success, the session is regenerated to prevent session fixation attacks. "
    "The user is then redirected to their role-specific dashboard.")
code_block(doc,"""public function login(Request $request)
{
    $validator = Validator::make($request->all(), [
        'email'    => ['required', 'email'],
        'password' => ['required'],
    ]);

    if ($validator->fails()) {
        return back()->withErrors($validator)
                     ->with('open_modal', 'login');
    }

    if (Auth::attempt($request->only('email', 'password'), true)) {
        $request->session()->regenerate();   // prevent session fixation
        return $this->redirectByRole();       // → correct dashboard by role
    }

    return back()->withErrors([
        'email' => 'These credentials do not match our records.',
    ]);
}

private function redirectByRole()
{
    $user = Auth::user();
    if ($user->isAdmin())    return redirect()->route('admin.dashboard');
    if ($user->isLandlord()) return redirect()->route('landlord.dashboard');
    return redirect()->route('tenant.dashboard');
}""")

doc.add_paragraph()
subheading(doc,"8.2  Registration  (AuthController::register)")
para(doc,
    "Registration enforces strong validation rules with user-friendly custom messages. "
    "Password strength is enforced using Laravel's Password rule object.")
code_block(doc,"""$validator = Validator::make($request->all(), [
    'name'     => ['required', 'string', 'max:100',
                   'regex:/^[\\pL\\s\\-\\\']+$/u'],   // letters, spaces, hyphens
    'email'    => ['required', 'email:rfc', 'max:150', 'unique:users,email'],
    'phone'    => ['nullable', 'regex:/^\\+?[0-9\\s\\-]{7,20}$/'],
    'role'     => ['required', 'in:landlord,tenant'],
    'password' => ['required', 'confirmed',
                   Password::min(8)->mixedCase()->numbers()->symbols()],
]);""")

doc.add_paragraph()
subheading(doc,"8.3  Session Lock Screen  (AuthController::showLock / unlock)")
para(doc,
    "When a user has a persistent session and reopens the browser, Auth::viaRemember() "
    "returns true and they are redirected to /lock before accessing any page. "
    "The unlock form requires their password and is rate-limited to 5 attempts "
    "with a 60-second decay window.")
code_block(doc,"""public function unlock(Request $request)
{
    $throttleKey = 'session-unlock:' . Auth::id();

    if (RateLimiter::tooManyAttempts($throttleKey, 5)) {
        $seconds = RateLimiter::availableIn($throttleKey);
        return back()->withErrors([
            'password' => "Too many failed attempts. Try again in {$seconds} seconds.",
        ]);
    }

    if (!Hash::check($request->password, Auth::user()->password)) {
        RateLimiter::hit($throttleKey, 60);   // 60-second decay
        $remaining = 5 - RateLimiter::attempts($throttleKey);
        return back()->withErrors([
            'password' => "Incorrect password. {$remaining} attempts remaining.",
        ]);
    }

    RateLimiter::clear($throttleKey);
    session()->forget('auth.locked');
    $request->session()->regenerate();
    return $this->redirectByRole();
}""")

insert_image(doc, os.path.join(PUB,"signIn","signin1.png"), width=5.5,
             caption="Figure 5 – Sign-In Page with Lock Screen Support")

doc.add_paragraph()
subheading(doc,"8.4  Force Password Change")
para(doc,
    "When a landlord creates a tenant account and sends an invitation, the tenant's "
    "force_password_change flag is set to true. The ForcePasswordChange middleware "
    "intercepts every tenant route and redirects to /tenant/change-password until "
    "the tenant sets a new password.")
code_block(doc,"""// ForcePasswordChange middleware
if ($user && $user->force_password_change) {
    if (! $request->routeIs('tenant.password.change',
                             'tenant.password.update', 'logout')) {
        return redirect()->route('tenant.password.change')
            ->with('warning', 'You must set a new password before continuing.');
    }
}

// TenantController::changePasswordUpdate
Auth::user()->update([
    'password'              => Hash::make($validated['password']),
    'force_password_change' => false,
]);""")

doc.add_paragraph()
subheading(doc,"8.5  Payment-Level Authorization")
para(doc,
    "Within PaymentController, every action verifies the requesting landlord owns "
    "the property associated with the payment, preventing one landlord from "
    "accessing another's payments.")
code_block(doc,"""private function authorizePayment(Payment $payment): void
{
    $propertyIds = Property::where('landlord_id', Auth::id())->pluck('id');

    $allowed = $payment->lease
        ? $propertyIds->contains($payment->lease->property_id)
        : false;

    abort_if(! $allowed, 403);   // → HTTP 403 Forbidden
}""")

# ── save Part 1 ──────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Part 1 saved: {OUT}")
